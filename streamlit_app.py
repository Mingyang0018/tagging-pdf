import streamlit as st
import fitz  # PyMuPDF
import json
import uuid
import re
import io
import pandas as pd
from streamlit_ace import st_ace
from docx import Document  # 用于生成 Word 文档

st.set_page_config(page_title="PDF Tagging Tool", layout="centered", initial_sidebar_state="expanded")
st.title("PDF Tagging Tool")

uploaded_file = st.file_uploader("Upload PDF File", type=["pdf"])

if uploaded_file:
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    pages = [page.get_text("text") for page in doc]

    # Add page number headers to each page text
    pages_with_numbers = [f"--- {i+1} ---\n{page}" for i, page in enumerate(pages)]
    joined_text = "\n\n".join(pages_with_numbers)

    # Editable ACE Editor
    edited = st_ace(
        value=joined_text,
        language="text",
        theme="chrome",
        height=800,
        key="ace-editor",
        auto_update=True
    )

    if edited:
        # Split and clean text page-wise
        edited_text = re.split(r'^--- \d+ ---', edited, flags=re.MULTILINE)[1:]
        edited_text = [text.strip() for text in edited_text if text.strip()]

        # JSON export
        export_data = {
            "document": uploaded_file.name,
            "edited_text": edited_text
        }
        json_data = json.dumps(export_data, indent=3, ensure_ascii=False)
        st.download_button(
            label="Download JSON",
            data=json_data,
            file_name=f"{uploaded_file.name}_{uuid.uuid4()}.json",
            mime="application/json"
        )

        # CSV export
        df = pd.DataFrame({
            "page_number": list(range(1, len(edited_text) + 1)),
            "text": edited_text
        })
        csv_buffer = io.StringIO()
        df.to_csv(csv_buffer, index=False)
        st.download_button(
            label="Download CSV",
            data=csv_buffer.getvalue(),
            file_name=f"{uploaded_file.name}_{uuid.uuid4()}.csv",
            mime="text/csv"
        )

        # DOCX export
        docx_buffer = io.BytesIO()
        docx_doc = Document()
        docx_doc.add_heading(f"{uploaded_file.name}", 0)

        for i, page_text in enumerate(edited_text):
            docx_doc.add_heading(f"Page {i+1}", level=1)
            docx_doc.add_paragraph(page_text)

        docx_doc.save(docx_buffer)
        docx_buffer.seek(0)

        st.download_button(
            label="Download DOCX",
            data=docx_buffer,
            file_name=f"{uploaded_file.name}_{uuid.uuid4()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
